"""
DOCX export functionality using python-docx.

This module provides DOCX generation from report data,
following functional programming principles with pure transformation functions.
"""

import io
import logging
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
else:
    DocxDocument = Any

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None  # type: ignore[assignment]

from src.models.report import Report, ReportFormat, ReportOutput, ReportSection
from src.services.report_config import ReportSettings

logger = logging.getLogger(__name__)


class DOCXExportError(Exception):
    """Exception raised during DOCX export."""
    pass


class DOCXExporter:
    """Service for exporting reports to DOCX format using python-docx."""

    def __init__(self, settings: ReportSettings | None = None):
        """Initialize DOCX exporter."""
        if not PYTHON_DOCX_AVAILABLE:
            raise DOCXExportError(
                "python-docx is not available. Install with: pip install python-docx"
            )

        self.settings = settings or ReportSettings()
    
    def export_to_docx(self, report: Report) -> ReportOutput:
        """
        Export report to DOCX format.
        
        Args:
            report: Report object to export
            
        Returns:
            ReportOutput with DOCX content
        """
        try:
            logger.info(f"Starting DOCX export for report: {report.id}")
            
            # Create document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add title page
            self._add_title_page(doc, report)
            
            # Add metadata
            self._add_metadata_section(doc, report)
            
            # Add executive summary
            if report.executive_summary and report.configuration.include_executive_summary:
                self._add_executive_summary(doc, report)
            
            # Add table of contents placeholder
            if report.configuration.include_toc:
                self._add_toc_placeholder(doc)
            
            # Add sections
            for section in report.sections:
                self._add_section(doc, section)
            
            # Add citations
            if report.citations and report.configuration.include_citations:
                self._add_citations(doc, report)
            
            # Save to bytes
            docx_bytes = self._document_to_bytes(doc)

            return ReportOutput(
                format=ReportFormat.DOCX,
                content=docx_bytes,
                file_path=None,
                file_size=len(docx_bytes),
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                encoding="binary"
            )
            
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            raise DOCXExportError(f"Failed to generate DOCX: {e}")
    
    def _setup_document_styles(self, doc: DocxDocument) -> None:
        """Set up custom document styles."""
        # Get default styles
        styles = doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Create custom styles
        try:
            # Title style
            if 'Report Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = 'Calibri'
                title_font.size = Pt(18)
                title_font.bold = True
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(12)
            
            # Heading styles are already available
            
            # Custom metadata style
            if 'Metadata' not in [s.name for s in styles]:
                metadata_style = styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
                metadata_font = metadata_style.font
                metadata_font.name = 'Calibri'
                metadata_font.size = Pt(10)
                metadata_font.italic = True
        
        except Exception as e:
            logger.warning(f"Could not create custom styles: {e}")
    
    def _add_title_page(self, doc: DocxDocument, report: Report) -> None:
        """Add title page to document."""
        # Title
        title_paragraph = doc.add_paragraph(report.title)
        title_paragraph.style = 'Report Title'
        
        # Subtitle
        if report.configuration.type:
            subtitle = f"{report.configuration.type.value.replace('_', ' ').title()} Report"
            subtitle_paragraph = doc.add_paragraph(subtitle)
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Research question
        doc.add_paragraph()  # Empty line
        question_paragraph = doc.add_paragraph("Research Question:")
        question_paragraph.runs[0].bold = True
        doc.add_paragraph(report.query)
        
        # Author and date
        doc.add_paragraph()  # Empty line
        if report.configuration.author_name:
            author_paragraph = doc.add_paragraph(f"Author: {report.configuration.author_name}")
            author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if report.configuration.institution:
            inst_paragraph = doc.add_paragraph(report.configuration.institution)
            inst_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        date_paragraph = doc.add_paragraph(f"Generated: {report.metadata.generated_at.strftime('%B %d, %Y')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        doc.add_page_break()  # type: ignore[no-untyped-call]
    
    def _add_metadata_section(self, doc: DocxDocument, report: Report) -> None:
        """Add metadata section to document."""
        # Metadata heading
        doc.add_heading('Report Metadata', level=1)
        
        # Create metadata table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Attribute'
        header_cells[1].text = 'Value'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add metadata rows
        metadata_items = [
            ('Quality Score', f"{report.metadata.quality_score * 100:.1f}%"),
            ('Confidence Score', f"{report.metadata.confidence_score * 100:.1f}%"),
            ('Sources Analyzed', str(report.metadata.total_sources)),
            ('Citations', str(report.metadata.total_citations)),
            ('Word Count', f"{report.metadata.word_count:,}"),
        ]
        
        if report.domains:
            metadata_items.append(('Research Domains', ', '.join(report.domains)))
            
        if report.metadata.agents_used:
            metadata_items.append(('AI Agents Used', ', '.join(report.metadata.agents_used)))
        
        for attribute, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = attribute
            row_cells[1].text = value
        
        doc.add_paragraph()  # Empty line
    
    def _add_executive_summary(self, doc: DocxDocument, report: Report) -> None:
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)

        # Convert markdown-like content to plain text
        summary_text = self._convert_markdown_to_text(report.executive_summary or "")
        doc.add_paragraph(summary_text)
        
        doc.add_paragraph()  # Empty line
    
    def _add_toc_placeholder(self, doc: DocxDocument) -> None:
        """Add table of contents placeholder."""
        doc.add_heading('Table of Contents', level=1)
        
        toc_paragraph = doc.add_paragraph(
            "Note: Table of contents can be generated in Microsoft Word using "
            "References > Table of Contents after opening this document."
        )
        toc_paragraph.style = 'Metadata'
        
        doc.add_paragraph()  # Empty line
    
    def _add_section(self, doc: DocxDocument, section: ReportSection) -> None:
        """Add a report section to the document."""
        # Section heading
        heading_level = min(section.level + 1, 6)  # Word supports up to 6 heading levels
        doc.add_heading(section.title, level=heading_level)
        
        # Section content
        content_text = self._convert_markdown_to_text(section.content)
        if content_text.strip():
            doc.add_paragraph(content_text)
        
        # Subsections
        for subsection in section.subsections:
            self._add_subsection(doc, subsection, section.level + 1)
        
        doc.add_paragraph()  # Empty line after section
    
    def _add_subsection(self, doc: DocxDocument, subsection: ReportSection, level: int) -> None:
        """Add a subsection to the document."""
        heading_level = min(level + 1, 6)
        doc.add_heading(subsection.title, level=heading_level)
        
        content_text = self._convert_markdown_to_text(subsection.content)
        if content_text.strip():
            doc.add_paragraph(content_text)
    
    def _add_citations(self, doc: DocxDocument, report: Report) -> None:
        """Add citations section to document."""
        doc.add_heading('References', level=1)
        
        for i, citation in enumerate(report.citations, 1):
            # Format citation
            citation_text = citation.format_citation(report.configuration.citation_style)
            
            # Add as numbered paragraph
            citation_paragraph = doc.add_paragraph(f"{i}. {citation_text}")
            citation_paragraph.style = 'Normal'
        
        doc.add_paragraph()  # Empty line
    
    def _convert_markdown_to_text(self, markdown_text: str) -> str:
        """Convert simple markdown to plain text for DOCX."""
        if not markdown_text:
            return ""
        
        # Remove markdown headers
        text = markdown_text
        lines = text.split('\n')
        clean_lines = []
        
        for line in lines:
            # Remove header markers
            if line.strip().startswith('#'):
                clean_line = line.strip().lstrip('#').strip()
                clean_lines.append(clean_line)
            # Remove bold/italic markers (basic)
            elif line.strip():
                clean_line = line.replace('**', '').replace('*', '').replace('`', '')
                clean_lines.append(clean_line)
            else:
                clean_lines.append('')
        
        return '\n'.join(clean_lines)
    
    def _document_to_bytes(self, doc: DocxDocument) -> bytes:
        """Convert Document object to bytes."""
        # Save document to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        return doc_bytes.getvalue()


def create_docx_exporter(settings: ReportSettings | None = None) -> DOCXExporter:
    """Factory function to create a DOCX exporter."""
    return DOCXExporter(settings)


__all__ = [
    "DOCXExportError",
    "DOCXExporter",
    "create_docx_exporter",
]